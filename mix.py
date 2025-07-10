import sys
import sqlite3
import os
import json
import speech_recognition as sr
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QPushButton, QTableWidget, 
    QTableWidgetItem, QLineEdit, QLabel, QDialog, QHBoxLayout, QFileDialog,
    QListWidget, QListWidgetItem, QSplitter, QScrollArea, QGridLayout, QWidget,
    QGroupBox, QToolButton, QMenu, QCheckBox
)
from PyQt6.QtCore import Qt
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from g4f.client import Client
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
try:
    from PyQt6.QtCharts import QChart, QChartView, QBarSeries, QBarSet, QValueAxis, QPieSeries, QLineSeries
    from PyQt6.QtCore import QPointF
    from PyQt6.QtGui import QPainter
    HAS_QTCHARTS = True
except ModuleNotFoundError:
    HAS_QTCHARTS = False

# Реєструємо шрифт Arial та Arial-Bold
pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/arial.ttf"))
pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))

class LoginDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Вхід")
        self.setGeometry(200, 200, 300, 120)
        layout = QVBoxLayout()

        # Поле для введення імені користувача
        self.user_label = QLabel("Користувач:")
        self.user_input = QLineEdit()
        layout.addWidget(self.user_label)
        layout.addWidget(self.user_input)

        # Поле для введення пароля
        self.pass_label = QLabel("Пароль:")
        self.pass_input = QLineEdit()
        self.pass_input.setEchoMode(QLineEdit.EchoMode.Password)
        layout.addWidget(self.pass_label)
        layout.addWidget(self.pass_input)

        # Кнопка для входу 
        self.login_button = QPushButton("Увійти")
        self.login_button.clicked.connect(self.attempt_login)
        layout.addWidget(self.login_button)

        self.setLayout(layout)

    def attempt_login(self):
        username = self.user_input.text()
        password = self.pass_input.text()
        
        # Перевірка у базі даних
        try:
            # Використовуємо окрему базу даних для аутентифікації або вашу основну
            conn = sqlite3.connect("user_management.db")  # або self.db_path
            cursor = conn.cursor()
            
            # Створити таблицю, якщо вона не існує
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                user_id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL,
                role TEXT DEFAULT 'user'
            )''')
            
            # Для першого запуску: додати адміністратора
            cursor.execute("SELECT COUNT(*) FROM users")
            if cursor.fetchone()[0] == 0:
                cursor.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", 
                              ("admin", "admin123", "admin"))
                conn.commit()
            
            # Перевірити користувача
            cursor.execute("SELECT user_id, role FROM users WHERE username=? AND password=?", 
                          (username, password))
            user_data = cursor.fetchone()
            conn.close()
            
            if user_data:
                self.user_id = user_data[0]
                self.user_role = user_data[1]
                self.accept()
                return
                
        except Exception as e:
            print(f"Помилка при вході: {e}")
        
        # Якщо невдалий вхід
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.warning(self, "Помилка входу", "Невірний логін або пароль")

class SQLApp(QWidget):
    def __init__(self, user_id=None, user_role=None):
        super().__init__()
        self.user_id = user_id
        self.user_role = user_role
        
        # === Window configuration ===
        self.setWindowTitle("SQL Viewer")
        self.setGeometry(100, 100, 1000, 700)
        
        # === Instance variables ===
        self.db_path = os.path.join(os.path.dirname(__file__), "project_management.db")
        self.pdf_export_path = os.path.expanduser("~/Documents")
        self.history = []
        self.favorites = []  # Change this to a dictionary to store both name and query
        self.favorites_data = {}  # Will store queries with their names as keys
        self.history_file = f"user_history_{self.user_id}.json" if self.user_id else "guest_history.json"
        
        # === UI initialization ===
        # Create main horizontal splitter
        main_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Create sidebar for history and favorites
        sidebar = QWidget()
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(0, 0, 0, 0)
        
        # History section
        history_group = QGroupBox("Історія запитів")
        history_layout = QVBoxLayout()
        
        # Add a layout for the history header with clear button
        history_header_layout = QHBoxLayout()
        history_header_label = QLabel("Історія")
        history_header_layout.addWidget(history_header_label)
        
        # Add clear history button
        self.clear_history_button = QPushButton("🗑️")
        self.clear_history_button.setToolTip("Очистити історію")
        self.clear_history_button.clicked.connect(self.clear_history)
        history_header_layout.addWidget(self.clear_history_button)
        
        history_layout.addLayout(history_header_layout)
        
        self.history_list = QListWidget()
        self.history_list.itemClicked.connect(self.load_from_history)
        history_layout.addWidget(self.history_list)
        history_group.setLayout(history_layout)
        sidebar_layout.addWidget(history_group)
        
        # Favorites section
        favorites_group = QGroupBox("Закріплені запити")
        favorites_layout = QVBoxLayout()

        # Add a layout for the favorites header with delete button
        favorites_header_layout = QHBoxLayout()
        favorites_header_label = QLabel("Закріплені")
        favorites_header_layout.addWidget(favorites_header_label)

        # Add delete favorite button
        self.delete_favorite_button = QPushButton("🗑️")
        self.delete_favorite_button.setToolTip("Видалити закріплений запит")
        self.delete_favorite_button.clicked.connect(self.delete_favorite)
        favorites_header_layout.addWidget(self.delete_favorite_button)

        favorites_layout.addLayout(favorites_header_layout)

        self.favorites_list = QListWidget()
        self.favorites_list.itemClicked.connect(self.load_from_favorites)
        favorites_layout.addWidget(self.favorites_list)
        favorites_group.setLayout(favorites_layout)
        sidebar_layout.addWidget(favorites_group)
        
        # Add sidebar to splitter
        main_splitter.addWidget(sidebar)
        
        # Main content area
        main_content = QWidget()
        layout = QVBoxLayout(main_content)
        
        # Database selection section
        db_layout = QHBoxLayout()
        self.db_label = QLabel("База даних:")
        db_layout.addWidget(self.db_label)
        
        self.db_path_label = QLabel(self.db_path)
        db_layout.addWidget(self.db_path_label, 1)  # Give this label more space
        
        self.db_select_button = QPushButton("Вибрати БД")
        self.db_select_button.clicked.connect(self.select_database)
        db_layout.addWidget(self.db_select_button)
        
        layout.addLayout(db_layout)

        # Tables access section with dropdown/grid for better handling of many tables
        self.tables_label = QLabel("Таблиці бази даних:")
        layout.addWidget(self.tables_label)
        
        # Create scrollable area for table buttons
        tables_scroll_area = QScrollArea()
        tables_scroll_area.setWidgetResizable(True)
        tables_container = QWidget()
        self.tables_grid = QGridLayout(tables_container)
        tables_scroll_area.setWidget(tables_container)
        tables_scroll_area.setMaximumHeight(120)  # Limit height
        layout.addWidget(tables_scroll_area)

        # Поле для введення запиту з голосовим введенням
        query_layout = QHBoxLayout()
        self.query_label = QLabel("Введіть ваш запит:")
        query_layout.addWidget(self.query_label)
        
        self.query_input = QLineEdit()
        query_layout.addWidget(self.query_input)
        
        # Додаємо кнопку голосового введення
        self.voice_button = QPushButton("🎤")
        self.voice_button.setToolTip("Голосове введення")
        self.voice_button.clicked.connect(self.voice_input)
        query_layout.addWidget(self.voice_button)
        
        # Add favorite button
        self.pin_button = QPushButton("📌")
        self.pin_button.setToolTip("Закріпити запит")
        self.pin_button.clicked.connect(self.add_to_favorites)
        query_layout.addWidget(self.pin_button)
        
        layout.addLayout(query_layout)

        # Основні кнопки для генерації та виконання - у новому окремому layout
        main_buttons_layout = QHBoxLayout()
        
        # Нова кнопка для використання AI (g4f) - тепер на першому місці
        self.ai_button = QPushButton("Генерувати ШІ")
        self.ai_button.clicked.connect(self.ai_query)
        self.ai_button.setMinimumHeight(40)  # Робимо кнопку більшою
        main_buttons_layout.addWidget(self.ai_button)
        
        # Кнопка виконання запиту - з новою назвою
        self.run_button = QPushButton("Виконати SQL запит")
        self.run_button.clicked.connect(self.execute_query)
        self.run_button.setMinimumHeight(40)  # Робимо кнопку більшою
        main_buttons_layout.addWidget(self.run_button)
        
        layout.addLayout(main_buttons_layout)

        # Додаткові кнопки для INSERT, DELETE, UPDATE
        btn_layout = QHBoxLayout()
        self.insert_button = QPushButton("Додати")
        self.insert_button.clicked.connect(self.insert_row)
        btn_layout.addWidget(self.insert_button)

        self.delete_button = QPushButton("Видалити")
        self.delete_button.clicked.connect(self.delete_row)
        btn_layout.addWidget(self.delete_button)

        self.update_button = QPushButton("Оновити")
        self.update_button.clicked.connect(self.update_row)
        btn_layout.addWidget(self.update_button)

        layout.addLayout(btn_layout)

        # Додаємо меню налаштувань та експорту
        settings_layout = QHBoxLayout()

        # Кнопка експорту в PDF
        self.export_pdf_button = QPushButton("📄 Експорт у PDF")
        self.export_pdf_button.setToolTip("Зберегти поточні результати у PDF файл")
        self.export_pdf_button.clicked.connect(self.export_to_pdf)
        settings_layout.addWidget(self.export_pdf_button)

        # Кнопка експорту у Word
        self.export_word_button = QPushButton("📝 Експорт у Word")
        self.export_word_button.setToolTip("Зберегти поточні результати у DOCX файл")
        self.export_word_button.clicked.connect(self.export_to_word)
        settings_layout.addWidget(self.export_word_button)

        # Кнопка експорту у Excel
        self.export_excel_button = QPushButton("📊 Експорт у Excel")
        self.export_excel_button.setToolTip("Зберегти поточні результати у XLSX файл")
        self.export_excel_button.clicked.connect(self.export_to_excel)
        settings_layout.addWidget(self.export_excel_button)

        layout.addLayout(settings_layout)
        
        # Таблиця для виводу результатів
        self.table = QTableWidget()
        layout.addWidget(self.table)

        # Add content to splitter and set up the final layout
        main_splitter.addWidget(main_content)
        
        # Set the ratio between sidebar and main content (1:3)
        main_splitter.setSizes([200, 800])
        
        # Set up the main layout with the splitter
        main_layout = QHBoxLayout(self)
        main_layout.addWidget(main_splitter)
        self.setLayout(main_layout)

        # Завантажуємо дані користувача
        self.load_user_preferences()
        self.load_user_favorites()
        self.load_user_history()
    
    # === USER DATA MANAGEMENT ===
    def load_user_preferences(self):
        """Завантаження налаштувань користувача"""
        if not self.user_id:
            return
            
        try:
            conn = sqlite3.connect("user_management.db")
            cursor = conn.cursor()
            
            # Створити таблицю, якщо вона не існує
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_preferences (
                user_id INTEGER PRIMARY KEY,
                last_db_path TEXT,
                pdf_export_path TEXT,
                FOREIGN KEY (user_id) REFERENCES users(user_id)
            )''')
            
            cursor.execute("SELECT last_db_path, pdf_export_path FROM user_preferences WHERE user_id=?", 
                         (self.user_id,))
            prefs = cursor.fetchone()
            conn.close()
            
            if prefs:
                db_path, pdf_path = prefs
                if db_path and os.path.exists(db_path):
                    self.db_path = db_path
                    self.db_path_label.setText(self.db_path)
                    self.load_database_tables()
                    
                if pdf_path:
                    self.pdf_export_path = pdf_path
                else:
                    self.pdf_export_path = os.path.expanduser("~/Documents")
            else:
                # Створити запис для нового користувача
                conn = sqlite3.connect("user_management.db")
                cursor = conn.cursor()
                cursor.execute("INSERT INTO user_preferences (user_id, last_db_path, pdf_export_path) VALUES (?, ?, ?)",
                             (self.user_id, self.db_path, os.path.expanduser("~/Documents")))
                conn.commit()
                conn.close()
                self.pdf_export_path = os.path.expanduser("~/Documents")
        except Exception as e:
            print(f"Помилка завантаження налаштувань: {e}")
    
    def load_user_favorites(self):
        """Завантаження закріплених запитів користувача"""
        if not self.user_id:
            return
            
        try:
            conn = sqlite3.connect("user_management.db")
            cursor = conn.cursor()
            
            # Створити таблицю, якщо вона не існує
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_favorites (
                favorite_id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                query TEXT NOT NULL,
                query_name TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(user_id)
            )''')
            
            cursor.execute("SELECT query, query_name FROM user_favorites WHERE user_id=? ORDER BY created_at DESC", 
                         (self.user_id,))
            favorites = cursor.fetchall()
            conn.close()
            
            # Очистити старі і завантажити нові
            self.favorites_list.clear()
            self.favorites = []
            self.favorites_data = {}  # Reset the dictionary
            
            for query, name in favorites:
                # Use name as display text, or truncated query if no name
                display_text = name if name else query[:50] + ("..." if len(query) > 50 else "")
                self.favorites.append(display_text)  # Store display text for UI
                self.favorites_data[display_text] = query  # Store full query with display text as key
                self.favorites_list.addItem(display_text)
                
        except Exception as e:
            print(f"Помилка завантаження закріплених запитів: {e}")

    def load_user_history(self):
        """Завантаження історії запитів користувача з JSON файлу"""
        try:
            history_path = os.path.join(os.path.dirname(__file__), self.history_file)
            if os.path.exists(history_path):
                with open(history_path, 'r', encoding='utf-8') as f:
                    self.history = json.load(f)
            else:
                self.history = []
                
            # Очистити старі і завантажити нові
            self.history_list.clear()
            
            for query in self.history:
                self.history_list.addItem(query[:50] + ("..." if len(query) > 50 else ""))
                
        except Exception as e:
            print(f"Помилка завантаження історії запитів: {e}")
            self.history = []

    def save_history_to_json(self):
        """Збереження історії запитів у JSON файл"""
        try:
            history_path = os.path.join(os.path.dirname(__file__), self.history_file)
            with open(history_path, 'w', encoding='utf-8') as f:
                json.dump(self.history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Помилка збереження історії запитів: {e}")

    def clear_history(self):
        """Очищення історії запитів"""
        from PyQt6.QtWidgets import QMessageBox
        
        reply = QMessageBox.question(self, 'Підтвердження',
                                     'Ви впевнені, що хочете очистити всю історію?',
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, 
                                     QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            # Очистити список та збережену історію
            self.history_list.clear()
            self.history = []
            self.save_history_to_json()
            QMessageBox.information(self, "Інформація", "Історію успішно очищено")

    def save_user_preferences(self):
        """Збереження налаштувань користувача"""
        if not self.user_id:
            return
            
        try:
            conn = sqlite3.connect("user_management.db")
            cursor = conn.cursor()
            cursor.execute("UPDATE user_preferences SET last_db_path=?, pdf_export_path=? WHERE user_id=?",
                         (self.db_path, self.pdf_export_path, self.user_id))
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Помилка збереження налаштувань: {e}")

    # === DATABASE OPERATIONS ===
    def select_database(self):
        """Відкриває діалог вибору файлу для вибору бази даних"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Виберіть файл бази даних SQLite",
            os.path.dirname(self.db_path),  # Start in current DB directory
            "SQLite Files (*.db *.sqlite *.db3);;All Files (*)"
        )
        
        if file_path:
            self.db_path = file_path
            self.db_path_label.setText(self.db_path)
            
            # Test connection and update quick access buttons
            try:
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()
                print(f"Connected to {self.db_path}")
                print(f"Tables found: {tables}")
                conn.close()
                
                # Update quick access table buttons
                self.update_table_buttons(tables)
            except Exception as e:
                print(f"Error connecting to database: {e}")

    def update_table_buttons(self, tables):
        """Creates buttons for quick access to each table in a grid layout"""
        # Clear existing buttons first
        for i in reversed(range(self.tables_grid.count())): 
            self.tables_grid.itemAt(i).widget().setParent(None)
        
        # Add new buttons for each table in a grid (4 columns)
        MAX_COLS = 4
        row, col = 0, 0
        
        for table in tables:
            table_name = table[0]
            # Skip sqlite_sequence as it's a system table
            if table_name == "sqlite_sequence":
                continue
                
            btn = QPushButton(table_name)
            btn.clicked.connect(lambda checked, name=table_name: self.show_table(name))
            self.tables_grid.addWidget(btn, row, col)
            
            # Move to next column or row
            col += 1
            if col >= MAX_COLS:
                col = 0
                row += 1

    def show_table(self, table_name):
        """Shows all data from the selected table"""
        query = f"SELECT * FROM {table_name}"
        self.query_input.setText(query)
        self.execute_query()

    def load_database_tables(self):
        """Завантаження таблиць з поточної бази даних"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            conn.close()
            
            # Update the table buttons UI
            self.update_table_buttons(tables)
        except Exception as e:
            print(f"Помилка при завантаженні таблиць: {e}")

    def execute_query(self):
        query = self.query_input.text()
        if not query:
            return
        
        # Перевірка прав доступу
        if not self.check_permissions(query):
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Помилка прав доступу", 
                              "У вас немає прав на виконання цього запиту")
            return
        
        # Додавання запиту в історію (змінено для JSON)
        if query not in self.history:
            self.history.insert(0, query)  # Додаємо на початок для хронологічного порядку
            if len(self.history) > 50:  # Обмежуємо кількість записів
                self.history = self.history[:50]
                
            # Оновлюємо UI
            self.history_list.clear()
            for q in self.history:
                self.history_list.addItem(q[:50] + ("..." if len(q) > 50 else ""))
                
            # Зберігаємо історію в JSON
            self.save_history_to_json()
        
        # Split the query into multiple statements
        queries = [q.strip() for q in query.split(';') if q.strip()]
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        results = []
        last_columns = []
        try:
            # Execute each query separately
            for q in queries:
                cursor.execute(q)
                try:
                    rows = cursor.fetchall()
                    if cursor.description:  # If this query returns results
                        columns = [desc[0] for desc in cursor.description]
                        results.extend(rows)
                        last_columns = columns
                except sqlite3.Error:
                    # This might be an INSERT/UPDATE/DELETE that doesn't return results
                    pass
            
            # Commit all changes at once
            conn.commit()
        except Exception as e:
            results = []
            last_columns = []
            print(f"Помилка запиту: {e}")
        finally:
            conn.close()

        # Clear the table
        self.table.clear()
        
        # Display results if any
        if last_columns:
            self.table.setColumnCount(len(last_columns))
            self.table.setHorizontalHeaderLabels(last_columns)
        else:
            self.table.setColumnCount(0)
        
        # Display results
        self.table.setRowCount(len(results))
        # Ховаємо номери рядків 
        self.table.setVerticalHeaderLabels([""] * len(results))
        
        for i, row in enumerate(results):
            for j, value in enumerate(row):
                self.table.setItem(i, j, QTableWidgetItem(str(value)))

    def check_permissions(self, query):
        """Перевірка прав доступу для виконання запиту"""
        if self.user_role == "admin":
            # Адміністратор має всі права
            return True
        
        if not self.user_id:
            # Якщо користувач не авторизований
            return False
        
        # Аналіз типу запиту
        query_lower = query.lower().strip()
        operation = None
        table_name = None
        
        if query_lower.startswith("select"):
            operation = "select"
        elif query_lower.startswith("insert"):
            operation = "insert"
        elif query_lower.startswith("update"):
            operation = "update"
        elif query_lower.startswith("delete"):
            operation = "delete"
        
        if operation:
            # Спрощений пошук назви таблиці
            # Це примітивний підхід, для складних запитів потрібен повноцінний SQL парсер
            if operation == "select":
                parts = query_lower.split("from")
                if len(parts) > 1:
                    table_part = parts[1].strip().split(" ")[0].split(";")[0]
                    table_name = table_part
            elif operation == "insert":
                parts = query_lower.split("into")
                if len(parts) > 1:
                    table_part = parts[1].strip().split(" ")[0].split("(")[0].split(";")[0]
                    table_name = table_part
            elif operation == "update":
                parts = query_lower.split("update")
                if len(parts) > 1:
                    table_part = parts[1].strip().split(" ")[0].split("set")[0].split(";")[0]
                    table_name = table_part
            elif operation == "delete":
                parts = query_lower.split("from")
                if len(parts) > 1:
                    table_part = parts[1].strip().split(" ")[0].split("where")[0].split(";")[0]
                    table_name = table_part
        
        if operation and table_name:
            # Перевірка прав доступу в базі
            try:
                conn = sqlite3.connect("user_management.db")
                cursor = conn.cursor()
                
                # Створити таблицю прав, якщо не існує
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS user_permissions (
                    permission_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER,
                    table_name TEXT NOT NULL,
                    can_select BOOLEAN DEFAULT 1,
                    can_insert BOOLEAN DEFAULT 0,
                    can_update BOOLEAN DEFAULT 0,
                    can_delete BOOLEAN DEFAULT 0,
                    FOREIGN KEY (user_id) REFERENCES users(user_id)
                )''')
                
                # Перевірка прав
                permission_field = f"can_{operation}"
                cursor.execute(f"SELECT {permission_field} FROM user_permissions WHERE user_id=? AND table_name=?", 
                             (self.user_id, table_name))
                result = cursor.fetchone()
                
                if not result:
                    # Якщо немає правила, створити дефолтне
                    defaults = {"can_select": 1, "can_insert": 0, "can_update": 0, "can_delete": 0}
                    cursor.execute(
                        "INSERT INTO user_permissions (user_id, table_name, can_select, can_insert, can_update, can_delete) VALUES (?, ?, ?, ?, ?, ?)", 
                        (self.user_id, table_name, defaults["can_select"], defaults["can_insert"], defaults["can_update"], defaults["can_delete"])
                    )
                    conn.commit()
                    conn.close()
                    return defaults.get(permission_field, 0) == 1
                
                conn.close()
                return result[0] == 1
            except Exception as e:
                print(f"Помилка перевірки прав доступу: {e}")
        
        # За замовчуванням дозволяємо SELECT і забороняємо інші операції
        if operation == "select":
            return True
        return False

    # === AI FEATURES ===
    def voice_input(self):
        """Функція для голосового введення запиту."""
        recognizer = sr.Recognizer()
        self.query_input.setText("Listening...")
        
        try:
            with sr.Microphone() as source:
                self.query_input.setText("Говоріть...")
                # Налаштування для шумного середовища
                recognizer.adjust_for_ambient_noise(source)
                audio = recognizer.listen(source, timeout=5)
                
            # Спроба розпізнати аудіо
            text = recognizer.recognize_google(audio, language='uk-UA')
            self.query_input.setText(text)
        except sr.UnknownValueError:
            self.query_input.setText("Не розпізнано, спробуйте ще раз")
        except sr.RequestError:
            self.query_input.setText("Помилка сервісу розпізнавання")
        except Exception as e:
            self.query_input.setText(f"Помилка: {str(e)}")

    def ai_query(self):
        query_text = self.query_input.text()
        
        # First validate the query
        is_valid, message = self.validate_ai_query(query_text)
        if not is_valid:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Валідація запиту", message)
            return
        
        # If valid, proceed with the regular AI query processing
        try:
            # 1) Connect and gather table info
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            schemas = {}
            for table in tables:
                table_name = table[0]
                cursor.execute(f"PRAGMA table_info({table_name})")
                columns = cursor.fetchall()
                schemas[table_name] = columns
            conn.close()

            # 2) Generate query from g4f
            if query_text.lower() == "exit":
                return
            
            # Show that processing is happening
            original_text = query_text
            self.query_input.setText("Генерую запит...")
            QApplication.processEvents()  # Update UI immediately
            
            # Process with updated mysql function
            sql_query = self.try_mysql(tables, schemas, query_text)
            print("Generated SQL:", sql_query)

            # 3) Execute and update table
            self.query_input.setText(sql_query)
            self.execute_query()
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Помилка ШІ", f"Помилка при генерації запиту: {str(e)}")
            self.query_input.setText(original_text)  # Restore original query

    def validate_ai_query(self, query_text):
        """Validate if a query is suitable for AI processing"""
        if not query_text or len(query_text.strip()) < 3:
            return False, "Запит занадто короткий. Будь ласка, введіть більш детальний запит."
        
        # Basic keyword check for SQL commands
        sql_keywords = ["select", "insert", "update", "delete", "create", "alter", "drop"]
        has_sql_command = any(keyword in query_text.lower() for keyword in sql_keywords)
        
        # If it's already a SQL command, we can use it directly
        if has_sql_command:
            return True, "SQL запит"
        
        # Check if it's likely a natural language query about data
        data_question_keywords = ["хто", "що", "де", "коли", "скільки", "як", "чому", "покажи", "знайди", "список", 
                                 "виведи", "дані", "інформація", "таблиця", "запис", "рядок"]
        has_question_indicators = any(keyword in query_text.lower() for keyword in data_question_keywords)
        
        if has_question_indicators:
            return True, "Запит природною мовою"
        
        # If it's very short and doesn't look like a question, validate with AI
        try:
            client = Client()
            response = client.chat.completions.create(
                model="gpt-4",  # Using a smaller model for quick validation
                messages=[
                    {
                        "role": "system",
                        "content": "Ти — програма, яка перевіряє чи є текст запитом до бази даних. " +
                                  "Якщо текст схожий на питання про дані або запит інформації — відповідай 'QUERY'. " +
                                  "Якщо це привітання, розмова або текст не пов'язаний з даними — відповідай 'NOT_QUERY'. " +
                                  "Відповідай лише одним словом: 'QUERY' або 'NOT_QUERY'."
                    },
                    {
                        "role": "user",
                        "content": query_text
                    }
                ],
                web_search=False
            )
            validation = response.choices[0].message.content.strip().upper()
            
            if "QUERY" in validation:
                return True, "Запит підтверджено ШІ"
            else:
                return False, "Ваш текст не схожий на запит до бази даних. Будь ласка, сформулюйте конкретний запит щодо даних."
                
        except Exception as e:
            print(f"Помилка валідації: {e}")
            # If validation fails, assume it's valid to avoid blocking legitimate queries
            return True, "Помилка перевірки запиту, продовжуємо виконання"

    def try_mysql(self, tables, schemas, query_text):
        """Generate SQL from natural language using AI."""
        try:
            # Prepare schema information for the AI
            db_schema = []
            for table in tables:
                table_name = table[0]
                if table_name == "sqlite_sequence":  # Skip system tables
                    continue
                    
                columns_info = []
                for col in schemas[table_name]:
                    # Format: (id, name, type, notnull, default_value, primary_key)
                    col_id, col_name, col_type, not_null, default_val, is_pk = col
                    col_desc = f"{col_name} ({col_type})"
                    if is_pk:
                        col_desc += " PRIMARY KEY"
                    columns_info.append(col_desc)
                    
                db_schema.append(f"Table: {table_name}\nColumns: {', '.join(columns_info)}")
            
            # Join all schema information
            schema_text = "\n\n".join(db_schema)
            
            # Check if schema is empty and handle it
            if not db_schema:
                return "SELECT 'No tables found in the database'"
            
            # Set up the prompt for the AI
            prompt = f"""Ти — програма, яка перетворює запити природною мовою на SQL-запити.

1. Виводь тільки чистий SQL-запит, без пояснень, лапок, форматування чи стилізації.
2. Якщо в запиті просять знайти щось за назвою або ім'ям — шукай як латиницею, так і кирилицею (використовуй LIKE і '%текст%', SQLite не підтримує ILIKE).
3. Твоя відповідь — лише один SQL-запит. Нічого більше.

СХЕМА БАЗИ ДАНИХ:
{schema_text}

ЗАПИТ КОРИСТУВАЧА: {query_text}

Поверни лише SQL-запит без додаткового тексту."""
            
            # Use g4f to generate SQL
            client = Client()
            response = client.chat.completions.create(
                model="gpt-4",  # Using a more powerful model for SQL generation
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert SQL developer. Generate only SQL queries without explanations or additional text."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                web_search=False
            )
            
            # Verify response is valid before accessing
            if not hasattr(response, 'choices') or not response.choices:
                debug_print("AI returned empty choices list")
                return "SELECT 'Error: AI returned empty response'"
            
            # Extract the SQL query from the response
            sql_query = response.choices[0].message.content.strip()
            
            # Basic cleaning of the response (remove markdown code blocks if present)
            if sql_query.startswith("```sql"):
                sql_query = sql_query[6:]
            if sql_query.startswith("```"):
                sql_query = sql_query[3:]
            if sql_query.endswith("```"):
                sql_query = sql_query[:-3]
                
            # Final cleanup
            sql_query = sql_query.strip()
            
            if not sql_query:
                return "SELECT 'Error: AI generated empty SQL query'"
            
            return sql_query
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            debug_print(f"Error in try_mysql: {str(e)}")
            # Fixed quote escaping with triple quotes
            return f"""SELECT 'Error generating SQL: {str(e).replace("'", "''")}'"""

    # === DATA MODIFICATION ===
    def insert_row(self):
        # Тут користувач може додавати свій INSERT-запит у self.query_input
        self.query_input.setText("INSERT INTO your_table (...) VALUES (...);")

    def delete_row(self):
        # Тут можна додавати DELETE-запит
        self.query_input.setText("DELETE FROM your_table WHERE id=...;")

    def update_row(self):
        # Тут можна додавати UPDATE-запит
        self.query_input.setText("UPDATE your_table SET column=value WHERE id=...;")

    # === HISTORY AND FAVORITES ===
    def load_from_history(self, item):
        """Load a query from history into the input field"""
        self.query_input.setText(item.text())
    
    def load_from_favorites(self, item):
        """Load a query from favorites into the input field and execute it"""
        display_text = item.text()
        if display_text in self.favorites_data:
            full_query = self.favorites_data[display_text]
            self.query_input.setText(full_query)
        else:
            self.query_input.setText(display_text)
        # Автоматично виконуємо запит
        self.execute_query()
    
    def add_to_favorites(self):
        """Add current query to favorites"""
        query = self.query_input.text()
        if not query:
            return
        
        # Check if this exact query is already in favorites
        if query in self.favorites_data.values():
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.information(self, "Дублікат", "Цей запит вже закріплено")
            return
            
        # Запит назви для закладки
        from PyQt6.QtWidgets import QInputDialog
        query_name, ok = QInputDialog.getText(self, "Назва закладки", "Введіть назву для закріпленого запиту:")
        
        if not ok:
            return
            
        # Generate display text
        display_text = query_name if query_name else query[:50] + ("..." if len(query) > 50 else "")
        
        # Додавання до локального UI та даних
        self.favorites.append(display_text)
        self.favorites_data[display_text] = query  # Store the full query with display text as key
        self.favorites_list.addItem(display_text)
        
        # Збереження у базу даних
        if self.user_id:
            try:
                conn = sqlite3.connect("user_management.db")
                cursor = conn.cursor()
                cursor.execute("INSERT INTO user_favorites (user_id, query, query_name) VALUES (?, ?, ?)",
                             (self.user_id, query, query_name))
                conn.commit()
                conn.close()
            except Exception as e:
                print(f"Помилка збереження закріпленого запиту: {e}")

    def delete_favorite(self):
        """Delete the selected favorite query"""
        current_item = self.favorites_list.currentItem()
        if not current_item:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Видалення закладки", "Спочатку виберіть закладку для видалення")
            return
        
        display_text = current_item.text()
        
        # Confirm deletion
        from PyQt6.QtWidgets import QMessageBox
        reply = QMessageBox.question(self, 'Підтвердження',
                                    f'Ви впевнені, що хочете видалити закладку "{display_text}"?',
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, 
                                    QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            # Remove from UI
            row = self.favorites_list.row(current_item)
            self.favorites_list.takeItem(row)
            
            # Remove from local data
            if display_text in self.favorites_data:
                query = self.favorites_data[display_text]
                if display_text in self.favorites:
                    self.favorites.remove(display_text)
                del self.favorites_data[display_text]
                
                # Remove from database
                if self.user_id:
                    try:
                        conn = sqlite3.connect("user_management.db")
                        cursor = conn.cursor()
                        # Try to find by query_name or by query itself
                        cursor.execute("DELETE FROM user_favorites WHERE user_id=? AND (query_name=? OR query=?)",
                                     (self.user_id, display_text, query))
                        conn.commit()
                        conn.close()
                        QMessageBox.information(self, "Видалення закладки", "Закладку успішно видалено")
                    except Exception as e:
                        print(f"Помилка видалення закладки: {e}")
                        QMessageBox.warning(self, "Помилка", f"Не вдалося видалити закладку: {str(e)}")

    # === EXPORT AND SETTINGS ===
    def generate_pdf_filename(self, query):
        """Generate a PDF filename based on the query content using AI"""
        try:
            # Спрощена версія - просто повернути базову назву
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_name = f"sql_export_{timestamp}.pdf"
            
            # Спроба генерації через ШІ
            print("Генерую назву файлу за допомогою ШІ...")
            client = Client()
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "Створи коротку описову назву файлу з суфіксом .pdf на основі SQL запиту. "
                                   "Назва має бути змістовна, українською мовою, без пробілів (використовуй _ замість пробілів). "
                                   "Наприклад, для запиту 'SELECT * FROM employees' поверни 'Список_працівників.pdf'"
                    },
                    {
                        "role": "user",
                        "content": f"Створи назву файлу для запиту: {query}"
                    }
                ],
                web_search=False
            )
            filename = response.choices[0].message.content.strip()
            print(f"ШІ згенерував назву: {filename}")
            
            # Ensure it has .pdf extension
            if not filename.lower().endswith('.pdf'):
                filename += '.pdf'
                
            # Replace spaces with underscores
            filename = filename.replace(' ', '_')
            
            # Remove any invalid filename characters
            invalid_chars = '<>:"/\\|?*'
            for char in invalid_chars:
                filename = filename.replace(char, '_')
                
            print(f"Остаточна назва файлу: {filename}")
            return filename
        except Exception as e:
            print(f"Помилка при генерації назви файлу: {e}")
            # Return a fallback filename with timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return f"sql_export_{timestamp}.pdf"

    def export_to_pdf(self):
        try:
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.pagesizes import A4, landscape, portrait
            from reportlab.lib.units import mm
            from PyQt6.QtWidgets import QFileDialog, QMessageBox
            from datetime import datetime
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase import pdfmetrics

            # Try to register Arial, fallback to Helvetica if not found
            arial_path = "C:/Windows/Fonts/arial.ttf"
            arial_bold_path = "C:/Windows/Fonts/arialbd.ttf"
            try:
                pdfmetrics.registerFont(TTFont("Arial", arial_path))
                pdfmetrics.registerFont(TTFont("Arial-Bold", arial_bold_path))
                font_name = "Arial"
                font_bold = "Arial-Bold"
            except Exception:
                font_name = "Helvetica"
                font_bold = "Helvetica-Bold"

            query = self.query_input.text()
            if not query:
                QMessageBox.warning(self, "PDF Export", "Немає даних для експорту!")
                return

            suggested_name = self.generate_pdf_filename(query)
            if not self.pdf_export_path or not os.path.exists(self.pdf_export_path):
                self.pdf_export_path = os.path.expanduser("~/Documents")
            suggested_path = os.path.join(self.pdf_export_path, suggested_name)

            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Зберегти як PDF",
                suggested_path,
                "PDF Files (*.pdf)"
            )
            if not filename:
                return

            new_path = os.path.dirname(filename)
            self.pdf_export_path = new_path
            self.save_user_preferences()

            col_count = self.table.columnCount()
            row_count = self.table.rowCount()
            if col_count == 0:
                QMessageBox.warning(self, "PDF Export", "Немає даних для експорту!")
                return

            headers = [self.table.horizontalHeaderItem(i).text() for i in range(col_count)]
            data = [headers]
            for row in range(row_count):
                row_data = []
                for col in range(col_count):
                    item = self.table.item(row, col)
                    row_data.append(item.text() if item else "")
                data.append(row_data)

            # --- PDF page size/orientation logic ---
            PAGE_MARGIN = 20 * mm
            max_width = 160  # max width per column in points
            # Estimate total table width
            col_widths = []
            for col in range(col_count):
                maxlen = max([len(str(data[row][col])) for row in range(len(data))])
                width = min(40 + maxlen * 6, max_width)
                col_widths.append(width)
            table_width = sum(col_widths)
            # If table too wide for portrait, use landscape
            if table_width > (A4[0] - 2 * PAGE_MARGIN):
                pagesize = landscape(A4)
                available_width = pagesize[0] - 2 * PAGE_MARGIN
            else:
                pagesize = portrait(A4)
                available_width = pagesize[0] - 2 * PAGE_MARGIN
            # Scale columns if needed
            scale = min(1.0, available_width / table_width)
            col_widths = [w * scale for w in col_widths]

            doc = SimpleDocTemplate(filename, pagesize=pagesize, rightMargin=PAGE_MARGIN, leftMargin=PAGE_MARGIN, topMargin=PAGE_MARGIN, bottomMargin=PAGE_MARGIN)
            elements = []

            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='QueryHeader', fontName=font_bold, fontSize=16, spaceAfter=10, textColor=colors.darkblue, alignment=1))
            styles.add(ParagraphStyle(name='NormalCustom', fontName=font_name, fontSize=11, spaceAfter=8))

            elements.append(Paragraph("Результати SQL-запиту", styles['QueryHeader']))
            elements.append(Paragraph(f"<b>SQL запит:</b> {query}", styles['NormalCustom']))
            elements.append(Paragraph(f"<b>Час створення:</b> {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}", styles['NormalCustom']))
            elements.append(Spacer(1, 10))

            pdf_table = Table(data, colWidths=col_widths, repeatRows=1)
            style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1976d2")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), font_bold),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 1), (-1, -1), font_name),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ])
            pdf_table.setStyle(style)
            elements.append(pdf_table)
            doc.build(elements)
            QMessageBox.information(self, "PDF Export", f"Файл успішно збережено як:\n{filename}")
        except Exception as e:
            print(f"Помилка при експорті в PDF: {e}")
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Помилка експорту", f"Не вдалося створити PDF файл.\nПомилка: {str(e)}")

    def export_to_word(self):
        from PyQt6.QtWidgets import QFileDialog, QMessageBox
        from datetime import datetime
        import docx
        from docx.shared import RGBColor
        
        query = self.query_input.text()
        if not query:
            QMessageBox.warning(self, "Word Export", "Немає даних для експорту!")
            return

        suggested_name = self.generate_pdf_filename(query).replace('.pdf', '.docx')
        if not self.pdf_export_path or not os.path.exists(self.pdf_export_path):
            self.pdf_export_path = os.path.expanduser("~/Documents")
        suggested_path = os.path.join(self.pdf_export_path, suggested_name)

        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Зберегти як Word",
            suggested_path,
            "Word Files (*.docx)"
        )
        if not filename:
            return

        try:
            doc = Document()
            doc.add_heading("Результати SQL-запиту", 0)
            doc.add_paragraph(f"SQL запит: {query}")
            doc.add_paragraph(f"Час створення: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            doc.add_paragraph("")  # Відступ

            col_count = self.table.columnCount()
            row_count = self.table.rowCount()
            if col_count == 0:
                QMessageBox.warning(self, "Word Export", "Немає даних для експорту!")
                return

            table = doc.add_table(rows=row_count + 1, cols=col_count)
            table.style = 'Table Grid'

            # Заголовки (без використання XML напряму)
            hdr_cells = table.rows[0].cells
            for i in range(col_count):
                cell = hdr_cells[i]
                cell.text = self.table.horizontalHeaderItem(i).text()
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

            # Дані
            for row in range(row_count):
                row_cells = table.rows[row + 1].cells
                for col in range(col_count):
                    item = self.table.item(row, col)
                    row_cells[col].text = item.text() if item else ""
                    for p in row_cells[col].paragraphs:
                        p.alignment = 1  # CENTER
                        for run in p.runs:
                            run.font.size = Pt(10)

            doc.save(filename)
            QMessageBox.information(self, "Word Export", f"Файл успішно збережено як:\n{filename}")
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Помилка експорту", f"Не вдалося створити Word файл.\nПомилка: {str(e)}")

    def export_to_excel(self):
        from PyQt6.QtWidgets import QFileDialog, QMessageBox
        from datetime import datetime
        from openpyxl.styles import PatternFill

        query = self.query_input.text()
        if not query:
            QMessageBox.warning(self, "Excel Export", "Немає даних для експорту!")
            return

        suggested_name = self.generate_pdf_filename(query).replace('.pdf', '.xlsx')
        if not self.pdf_export_path or not os.path.exists(self.pdf_export_path):
            self.pdf_export_path = os.path.expanduser("~/Documents")
        suggested_path = os.path.join(self.pdf_export_path, suggested_name)

        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Зберегти як Excel",
            suggested_path,
            "Excel Files (*.xlsx)"
        )
        if not filename:
            return

        wb = Workbook()
        ws = wb.active
        ws.title = "SQL Results"

        col_count = self.table.columnCount()
        row_count = self.table.rowCount()
        if col_count == 0:
            QMessageBox.warning(self, "Excel Export", "Немає даних для експорту!")
            return

        # Заголовки
        headers = [self.table.horizontalHeaderItem(i).text() for i in range(col_count)]
        ws.append(headers)
        header_fill = PatternFill(start_color="1976d2", end_color="1976d2", fill_type="solid")
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.fill = header_fill

        # Дані
        for row in range(row_count):
            row_data = []
            for col in range(col_count):
                item = self.table.item(row, col)
                row_data.append(item.text() if item else "")
            ws.append(row_data)

        # Автоширина колонок
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = max(10, min(max_length + 2, 40))

        try:
            wb.save(filename)
            QMessageBox.information(self, "Excel Export", f"Файл успішно збережено як:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Помилка експорту", f"Не вдалося створити Excel файл.\nПомилка: {str(e)}")

# === Add these debug functions at the bottom of the file, before the if __name__ == "__main__": block ===
def debug_print(message):
    """Helper function to print debugging messages with timestamps"""
    from datetime import datetime
    timestamp = datetime.now().strftime("%H:%M:%S.%f")[:-3]
    print(f"[DEBUG {timestamp}] {message}")

def initialize_database():
    """Initialize necessary databases for the application"""
    # Create user_management.db for authentication
    try:
        conn = sqlite3.connect("user_management.db")
        cursor = conn.cursor()
        
        # Create users table if it doesn't exist
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            role TEXT DEFAULT 'user'
        )''')
        
        # Create default admin user if no users exist
        cursor.execute("SELECT COUNT(*) FROM users")
        if cursor.fetchone()[0] == 0:
            cursor.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", 
                          ("admin", "admin123", "admin"))
        
        # Create user_preferences table
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS user_preferences (
            user_id INTEGER PRIMARY KEY,
            last_db_path TEXT,
            pdf_export_path TEXT,
            FOREIGN KEY (user_id) REFERENCES users(user_id)
        )''')
        
        # Create user_permissions table
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS user_permissions (
            permission_id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            table_name TEXT NOT NULL,
            can_select BOOLEAN DEFAULT 1,
            can_insert BOOLEAN DEFAULT 0,
            can_update BOOLEAN DEFAULT 0,
            can_delete BOOLEAN DEFAULT 0,
            FOREIGN KEY (user_id) REFERENCES users(user_id)
        )''')
        
        # Create user_favorites table
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS user_favorites (
            favorite_id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            query TEXT NOT NULL,
            query_name TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(user_id)
        )''')
        
        conn.commit()
        conn.close()
        debug_print("User management database initialized successfully")
        
        # Ensure project_management.db exists (create an empty one if not present)
        project_db_path = os.path.join(os.path.dirname(__file__), "project_management.db")
        if not os.path.exists(project_db_path):
            conn = sqlite3.connect(project_db_path)
            conn.close()
            debug_print(f"Created empty project database at {project_db_path}")
            
    except Exception as e:
        debug_print(f"Error initializing database: {e}")
        raise

# Update the main execution block with debug statements
if __name__ == "__main__":
    try:
        debug_print("Starting application")
        app = QApplication(sys.argv)
        
        debug_print("Initializing database")
        initialize_database()
        
        debug_print("Creating login dialog")
        login_dialog = LoginDialog()
        debug_print("Showing login dialog")
        result = login_dialog.exec()
        debug_print(f"Login dialog result: {result}")
        
        if result == QDialog.DialogCode.Accepted:
            debug_print("Login successful, creating main window")
            window = SQLApp(login_dialog.user_id, login_dialog.user_role)
            debug_print("Showing main window")
            window.show()
            debug_print("Starting event loop")
            sys.exit(app.exec())
        else:
            debug_print("Login canceled or failed")
            sys.exit(0)
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


